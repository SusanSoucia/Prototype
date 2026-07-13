#!/usr/bin/env python3
"""
企业RAG 详细设计说明书 DOCX生成器 (完整版)
数据驱动架构 - 所有模块配置集中定义
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

DIAGRAMS = "/home/susan/MySpace/codeSpace/shixun/diagrams"
DOC = Document()

# ═══════════════════ STYLE SETUP ═══════════════════
style = DOC.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for i in range(1, 4):
    hs = DOC.styles[f'Heading {i}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.size = Pt([18, 15, 13][i - 1])

# ═══════════════════ HELPERS ═══════════════════
def P(text='', bold=False, size=12, align=None, indent=None, heading=None):
    """Add paragraph. If heading is set, use heading style."""
    if heading:
        DOC.add_heading(text, level=heading)
        return
    p = DOC.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.left_indent = Cm(indent)

def T(headers, rows, col_widths=None):
    """Add formatted table."""
    t = DOC.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); r.font.name = '宋体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4" w:val="clear"/>')
        c._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = t.rows[1 + ri].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10); r.font.name = '宋体'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if ri % 2 == 1:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                c._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    DOC.add_paragraph()
    return t

def IMG(name, w=5.5):
    """Add image."""
    path = os.path.join(DIAGRAMS, name)
    if os.path.exists(path):
        p = DOC.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(w))
        DOC.add_paragraph()
    else:
        P(f'[图片缺失: {name}]')

def io_table(headers, items, widths):
    """Input/output item table."""
    if not items or not isinstance(items[0], dict): return P(str(items))
    return T(headers, [[i.get(k,'') for k in ['name','type','source','desc']] for i in items], widths)

# ═══════════════════ MODULE SECTION GENERATOR ═══════════════════
def module(seq, mid, name, d):
    """Generate complete module section per template."""
    DOC.add_heading(f'{seq}．模块{seq-2}（{mid}）设计说明', level=1)

    sections = [
        ('模块描述', d.get('desc', f'{name}的详细设计说明。')),
        ('功能', d.get('func', '')),
        ('性能', d.get('perf', '满足系统整体性能指标要求。')),
        ('输入项', d.get('input', '')),
        ('输出项', d.get('output', '')),
        ('设计方法（算法）', d.get('algo', '本模块采用标准设计方法实现。')),
        ('流程逻辑', d.get('flow', '')),
        ('接口', d.get('iface', '')),
        ('存储分配', d.get('storage', '本模块使用系统统一分配的存储资源。')),
        ('注释设计', d.get('comments', '所有公开方法均包含Javadoc注释，说明参数、返回值和异常。')),
        ('限制条件', d.get('constraints', '无特殊限制条件。')),
        ('测试计划', d.get('test', '')),
        ('尚未解决的问题', d.get('unsolved', '无。')),
    ]

    for i, (title, content) in enumerate(sections):
        sub_num = f'{seq}.{i+1}'
        DOC.add_heading(f'{sub_num} {title}', level=2)

        if title == '功能' and isinstance(content, list):
            for f in content: P(f'• {f}')
        elif title == '输入项' and isinstance(content, list):
            io_table(['输入项', '类型/格式', '来源', '说明'], content, [3,3,3,6])
        elif title == '输出项' and isinstance(content, list):
            io_table(['输出项', '类型/格式', '目标', '说明'], content, [3,3,3,6])
        elif title == '接口' and isinstance(content, list) and content and isinstance(content[0], dict):
            T(['接口名称', '类型', '调用方式', '说明'],
              [[c.get(k,'') for k in ['name','type','method','desc']] for c in content], [3,2.5,3.5,6])
        elif title == '限制条件' and isinstance(content, list):
            for c in content: P(f'• {c}')
        elif title == '测试计划' and isinstance(content, list) and content and isinstance(content[0], dict):
            T(['测试项', '测试方法', '输入数据', '预期结果'],
              [[t.get(k,'') for k in ['item','method','input','expected']] for t in content], [3,3,4.5,4.5])
        elif title == '流程逻辑':
            P(content)
            if d.get('flow_img'): IMG(d['flow_img'])
        else:
            P(str(content))

    DOC.add_page_break()

# ═══════════════════ MODULE DATA ═══════════════════
# Each entry: (seq_number, module_id, module_name, data_dict)
MODULES = []

def reg(seq, mid, name, **kw):
    MODULES.append((seq, mid, name, kw))

# --- M1: Spring Boot ---
reg(3, 'M1', 'Spring Boot应用服务框架',
    desc='M1模块是整个系统的核心编排引擎，基于Spring Boot 3.4.2构建。它负责接收前端HTTP/WebSocket请求，编排所有其他模块的调用，实现ReAct Agent循环、混合搜索、文档处理、用户认证和权限管理等核心业务逻辑。M1是系统中唯一直接面向用户的模块，所有用户请求都必须经过M1进行路由、鉴权和编排。',
    func=[
        'HTTP REST API服务：提供全部后端接口（用户管理、知识库管理、搜索、对话、系统管理）',
        'WebSocket长连接管理：处理实时问答通信，支持流式token推送和用户中断生成',
        'ReAct Agent编排：实现LLM推理+工具调用的循环逻辑（最多4轮循环，最多8次工具调用）',
        '混合搜索编排：协调Embedding API和ES完成KNN+BM25双路检索与RRF融合',
        '文档上传管理：大文件分片上传、MinIO合并、Kafka异步任务分发',
        'JWT认证与授权：集成Spring Security，实现基于角色（USER/ADMIN）的访问控制',
        '组织标签多租户：三层可见性控制（个人userId/组织orgTag/公开public）',
        '速率限制：基于Redis滑动窗口算法，对问答和API调用进行频率控制',
        'Token限额管理：记录和扣减用户Embedding和LLM消耗，支持管理员追加配额',
    ],
    perf='• 并发WebSocket连接：支持50+并发长连接\n• HTTP接口响应时间：简单查询<50ms，搜索<500ms\n• WebSocket首token延迟：<200ms\n• 文件分片上传：每片5MB，并行传输\n• JVM堆内存：建议-Xmx2g -Xms1g\n• 线程池：Tomcat默认200线程，Kafka Consumer自定义线程池',
    input=[
        {'name': 'HTTP REST请求', 'type': 'JSON/Multipart', 'source': 'Vue 3前端', 'desc': '所有管理类API调用，携带JWT认证头'},
        {'name': 'WebSocket消息帧', 'type': 'Text JSON帧', 'source': '浏览器', 'desc': '用户问题输入、停止生成信号、心跳ping'},
        {'name': 'JWT Token', 'type': 'Bearer Token', 'source': 'HTTP Header / URL Query参数', 'desc': '包含userId、role、orgTags等用户身份信息'},
    ],
    output=[
        {'name': 'HTTP REST响应', 'type': 'JSON', 'source': 'Vue 3前端', 'desc': '统一{code, message, data}格式，code=0表示成功'},
        {'name': 'WebSocket流式推送', 'type': 'JSON帧序列', 'source': '浏览器', 'desc': '逐token的chunk帧、tool_call状态帧、completion完成帧、referenceMappings引用映射'},
    ],
    algo='【ReAct Agent循环算法】\n'
         '1. 构建系统Prompt，包含4个工具的函数定义（JSON Schema格式）\n'
         '2. 组装messages数组：[system_prompt, ...历史对话, user_question]\n'
         '3. 调用M2(DeepSeek) POST /v1/chat/completions，stream=true，携带tools参数\n'
         '4. 解析SSE事件流：\n'
         '   - delta.content → 封装为chunk帧，流式推送到浏览器\n'
         '   - delta.tool_calls → 收集完整tool_call后，分发到AgentToolRegistry执行\n'
         '5. 工具执行结果以role=tool消息追加到messages数组\n'
         '6. 回到步骤3，继续下一轮推理\n'
         '7. 终止条件：(a)LLM返回stop不调用工具 (b)达到4轮上限 (c)用户发送停止信号 (d)超时120秒\n\n'
         '【混合搜索编排算法】\n'
         '1. 调用M3(Embedding)将查询文本向量化为float[2048]\n'
         '2. 调用M4(ES)：KNN向量召回(recallK=topK×30) + BM25关键词过滤 + rescore重排序\n'
         '3. M4返回结果后，调用M5(MySQL)查询fileMd5→fileName映射，附加到搜索结果\n'
         '4. 权限过滤后返回最终topK结果',
    flow='用户问答场景的完整数据流（ReAct Agent循环+混合检索+流式推送），详见下图：',
    flow_img='02_qa_flow.png',
    iface=[
        {'name': 'ChatHandler.processMessage()', 'type': '内部Service', 'method': 'WebSocket消息入口', 'desc': '接收用户问题，启动ReAct循环，管理会话生命周期'},
        {'name': 'HybridSearchService.search()', 'type': '内部Service', 'method': 'Service层调用', 'desc': '协调Embedding+ES完成KNN+BM25混合检索'},
        {'name': 'AgentToolRegistry.execute()', 'type': '内部Service', 'method': '工具调用路由分发', 'desc': '根据toolName路由到search_knowledge/generate_summary/submit_feedback/knowledge_stats'},
        {'name': 'UploadController', 'type': 'REST Controller', 'method': 'POST /api/v1/upload/*', 'desc': '文件分片上传和合并触发端点'},
        {'name': 'ChatWebSocketHandler', 'type': 'WebSocket Handler', 'method': 'ws://host/chat/{token}', 'desc': 'WebSocket握手、认证、消息路由、连接注册'},
    ],
    storage='• M1自身不直接存储业务数据，所有持久化委托给M5(MySQL)和M6(Redis)\n• JVM堆内存暂存：WebSocket Session注册表、ReAct循环上下文、上传进度状态\n• 文件分片暂存：委托给M8(MinIO) chunks/{md5}/{idx}路径\n• 异步任务消息：委托给M7(Kafka) topic持久化',
    constraints=[
        '• 必须运行在Java 17及以上版本，依赖Spring Boot 3.4.x自动配置',
        '• WebSocket最大并发连接数受限于Tomcat线程池（默认200）',
        '• ReAct循环硬限制：最多4轮推理，最多8次工具调用，超时120秒',
        '• 单文件上传最大100MB（受MinIO分片数量和上传超时限制）',
        '• CORS仅允许配置的Origin（开发环境localhost:9527，生产环境同域）',
        '• 速率限制器依赖M6(Redis)，Redis不可用时降级为无限制模式',
    ],
    test=[
        {'item': 'WebSocket连接与认证', 'method': '前端集成测试', 'input': '有效JWT Token通过URL参数传递', 'expected': 'HTTP 101升级成功，收到{"type":"connection"}帧'},
        {'item': 'ReAct问答完整流程', 'method': 'Mock M2/M3/M4', 'input': '"企业2025年差旅标准是什么？"', 'expected': '流式返回含引用标记的回答，引用映射完整'},
        {'item': '用户中途停止生成', 'method': '发送STOP JSON帧', 'input': '在生成过程中发送停止信号', 'expected': 'SSE连接被中断，返回{"type":"completion","status":"cancelled"}'},
        {'item': '速率限制触发', 'method': 'JMeter压测', 'input': '1分钟内连续发送超过限制次数的请求', 'expected': '返回429 Too Many Requests错误'},
        {'item': 'JWT过期自动处理', 'method': '使用过期Token', 'input': '已过期的JWT Token', 'expected': 'WebSocket握手阶段返回401，前端弹出登录框'},
    ],
)

# --- M2: DeepSeek API ---
reg(4, 'M2', 'DeepSeek API大语言模型服务',
    desc='M2模块是系统的AI推理核心，通过DeepSeek云服务API提供大语言模型能力。在企业RAG中，M2承担双重角色：(1)作为ReAct Agent的"决策大脑"，根据用户问题自主判断是否需要调用工具（search_knowledge等）；(2)作为"生成引擎"，基于RAG检索到的文档片段生成带来源引用的自然语言回答。M2以OpenAI兼容的REST API形式提供，M1通过HTTP SSE流式接收生成结果，实现逐token推送的打字机效果。',
    func=[
        'ReAct推理与决策：分析用户问题，自主决定调用哪个工具、传什么参数',
        '知识增强生成（RAG Generation）：基于检索到的文档片段生成准确、有引用依据的回答',
        '摘要生成（Summarization）：通过generate_summary工具对多篇文档进行归纳总结',
        '流式内容输出（Streaming）：通过SSE协议逐token返回生成内容，支持前端打字机效果',
        '多轮对话上下文保持：在单次会话中追踪多轮对话的上下文关联',
        '多Provider可切换：通过ModelProvider路由机制支持切换不同的LLM服务商',
    ],
    perf='• 首token延迟：200-500ms（取决于模型负载）\n• 生成速度：30-50 tokens/秒\n• 上下文窗口：最大64K tokens\n• API并发限制：取决于API Key套餐（免费版约10次/分钟）\n• 服务可用性：99.9%（DeepSeek云服务SLA）',
    input=[
        {'name': 'Chat Completion请求', 'type': 'OpenAI格式JSON', 'source': 'M1(DeepSeekClient)', 'desc': 'messages数组（system+history+user）+ tools工具定义 + model=deepseek-chat + stream=true'},
    ],
    output=[
        {'name': 'SSE事件流', 'type': 'text/event-stream', 'source': 'M1(DeepSeekClient)', 'desc': '逐token的delta.content文本块 或 delta.tool_calls工具调用参数块，以finish_reason结束'},
    ],
    algo='M2对调用方（M1）是黑盒云服务，内部使用DeepSeek自研MoE（Mixture of Experts）大模型架构。\n\n'
         'M1调用参数规范：\n'
         '• model: "deepseek-chat"\n'
         '• messages: [{role:"system", content:系统Prompt}, {role:"user", content:用户问题}, ...]\n'
         '• tools: [4个工具的JSON Schema定义]\n'
         '• tool_choice: "auto"（LLM自主判断是否需要调用工具，无需人工预设）\n'
         '• stream: true（启用SSE流式输出模式）\n'
         '• temperature: 0.7（控制生成随机性，0=确定，1=最大随机）\n'
         '• max_tokens: 4096（单次生成最大token数）\n\n'
         'M1解析SSE事件流的处理步骤：\n'
         '1. SSE事件行格式：data: {"id":"...","choices":[{"delta":{"content":"年"},"index":0}]}\n'
         '2. 若delta含tool_calls字段 → 累积完整工具调用参数 → M1执行工具 → 结果追加到messages\n'
         '3. 若finish_reason="tool_calls" → 继续下一轮ReAct循环\n'
         '4. 若finish_reason="stop" → 本轮问答完成，保存对话记录',
    flow='M2在ReAct循环中充当推理和生成双重角色。用户问题首先进入M2进行决策（是否需要检索），检索完成后再次进入M2进行回答生成。详见M1模块流程图。',
    flow_img='02_qa_flow.png',
    iface=[
        {'name': 'POST /v1/chat/completions', 'type': 'REST API (OpenAI兼容)', 'method': 'HTTP POST + SSE流式响应', 'desc': 'DeepSeek对话补全API，支持流式和非流式模式'},
        {'name': 'DeepSeekClient.chat()', 'type': 'Java封装方法', 'method': 'OkHttp + EventSource', 'desc': 'M1封装的调用方法，处理API Key认证、请求构造、SSE事件解析和重试逻辑'},
    ],
    storage='M2是外部云服务，不涉及本地存储。API Key通过环境变量DEEPSEEK_API_KEY配置，支持多个Key轮换。',
    constraints=[
        '• 强依赖DeepSeek API服务可用性，服务中断时系统核心AI功能不可用',
        '• 需要有效的API Key和充足账户余额',
        '• 上下文窗口最大64K tokens，超出需按"先进先出"截断历史对话',
        '• 免费版API限制约10次/分钟，生产部署需升级套餐',
        '• SSE连接超时120秒（与ReAct循环超时保持一致）',
        '• 网络延迟直接影响用户感知的首token时间',
    ],
    test=[
        {'item': '基本知识问答', 'method': 'API直调测试', 'input': '"什么是RAG技术？"', 'expected': '返回关于检索增强生成的介绍性回答'},
        {'item': '工具调用决策能力', 'method': '携带tools参数测试', 'input': '"知识库中有多少份文档？"', 'expected': 'LLM返回tool_calls，调用knowledge_stats工具'},
        {'item': '流式输出完整性', 'method': '收集全部SSE chunk', 'input': '任意知识性问题', 'expected': '拼接后内容完整无截断，finish_reason=stop'},
        {'item': 'API Key失效处理', 'method': '使用已吊销的Key', 'input': '无效API Key', 'expected': 'M2返回401，M1记录错误日志并向前端返回服务不可用'},
    ],
)

# --- M3: Embedding API ---
reg(5, 'M3', 'DashScope Embedding向量化服务',
    desc='M3模块负责将文本转换为固定维度的浮点数向量（embedding），是实现语义检索的数学基础。企业RAG使用阿里云DashScope的text-embedding-v4模型，每条文本输出2048维归一化向量。M3在两类核心场景中被频繁调用：(1)文档入库时，将切分好的文本块批量向量化后存入ES索引；(2)用户问答时，将查询问题实时向量化后用于ES KNN检索。',
    func=[
        '文本向量化：将任意长度文本（≤8192 tokens）转换为float[2048]语义向量',
        '批量向量化：支持单次请求处理最多10条文本，大幅降低API调用次数，提升文档入库吞吐量',
        '查询向量化：将用户自然语言问题实时转换为查询向量，延迟控制在300ms以内',
        '向量归一化：输出向量已L2归一化（模长为1），可直接用于余弦相似度计算，无需二次处理',
    ],
    perf='• 单条文本处理延迟：100-300ms\n• 批量处理（10条/批）：500-1000ms\n• 向量维度：2048维float32\n• 支持语言：中文为主，英文兼容\n• Token消耗：约1 Token/中文字符，英文约1.3 Token/词\n• API并发限制：取决于DashScope套餐配额',
    input=[
        {'name': '文本列表', 'type': 'List<String> (1-10条)', 'source': 'M1(VectorizationService → EmbeddingClient)', 'desc': '待向量化的文本，单条查询或批量入库'},
    ],
    output=[
        {'name': '向量列表', 'type': 'List<float[2048]>', 'source': 'M1 → M4(ES索引/查询)', 'desc': '2048维归一化float32向量，已L2归一化'},
    ],
    algo='M3内部使用阿里云自研text-embedding-v4模型，对调用方是黑盒API。\n\n'
         'M1调用封装（EmbeddingClient.java）：\n'
         '1. HTTP POST请求：\n'
         '   URL: https://dashscope.aliyuncs.com/api/v1/services/embeddings/text-embedding/text-embedding\n'
         '   Headers: Authorization: Bearer {EMBEDDING_API_KEY}\n'
         '   Body: {"model":"text-embedding-v4","input":{"texts":["文本1","文本2",...]}}\n'
         '2. 响应解析：提取 output.embeddings[].embedding[] 数组\n'
         '3. 重试策略（指数退避）：网络异常自动重试3次，间隔1s→2s→4s\n'
         '4. 批量调度（VectorizationService）：\n'
         '   读取M5(MySQL) document_vectors表中待处理的文本分块\n'
         '   按batchSize=10分组，逐批调M3\n'
         '   将获得的向量写入M4(ES) knowledge_base索引',
    flow='M3在文档入库（异步批处理）和问答搜索（同步实时）两个场景中都被调用。文档入库流程详见M7模块流程图，问答搜索流程详见M1模块流程图。',
    iface=[
        {'name': 'EmbeddingClient.embed(String text)', 'type': 'Java方法', 'method': 'HTTP POST同步调用', 'desc': '单条文本向量化，用于实时查询场景'},
        {'name': 'EmbeddingClient.embedBatch(List<String> texts)', 'type': 'Java方法', 'method': 'HTTP POST同步调用', 'desc': '批量文本向量化（最多10条），用于文档入库场景'},
        {'name': 'POST /api/v1/services/embeddings/...', 'type': 'REST API (DashScope原生)', 'method': 'HTTPS POST', 'desc': '阿里云DashScope Embedding API端点'},
    ],
    storage='M3是外部云服务，不涉及本地存储。API Key通过环境变量EMBEDDING_API_KEY配置。向量本身存储在M4(ES)中。',
    constraints=[
        '• 强依赖DashScope API服务可用性，服务中断时搜索功能降级为纯BM25关键词检索',
        '• 单次批量请求最多10条文本，超出需自行分组',
        '• 文本单条最大长度8192 tokens（约6000中文字），超出需在调用前截断',
        '• 向量维度固定2048维，无法按需调整（与ES索引mapping绑定）',
        '• API调用消耗Token配额，文档入库时Embedding是主要成本来源',
        '• 网络延迟影响实时查询的搜索响应时间（目标<500ms）',
    ],
    test=[
        {'item': '单条文本向量化', 'method': 'API调用+维度验证', 'input': '"企业知识库RAG问答系统"', 'expected': '返回float[2048]，向量非全零，L2范数≈1.0'},
        {'item': '批量向量化（10条）', 'method': 'API调用+批量对比', 'input': '10条不同中英文混合文本', 'expected': '返回10个不同的float[2048]向量，两两余弦相似度<1'},
        {'item': '语义相似性验证', 'method': '余弦相似度计算', 'input': '"苹果"和"水果"的向量 vs "苹果"和"汽车"的向量', 'expected': '("苹果","水果")>0.7, ("苹果","汽车")<0.3'},
        {'item': '网络故障重试', 'method': 'Mock网络超时', 'input': '模拟connect timeout', 'expected': '3次指数退避重试后抛出自定义异常'},
    ],
)

# --- M4: Elasticsearch ---
reg(6, 'M4', 'Elasticsearch检索引擎',
    desc='M4模块是系统的核心检索引擎，承担向量检索、全文搜索和混合排序三大职责。企业RAG使用Elasticsearch 8.10版本，利用其dense_vector字段类型和KNN插件实现向量检索能力，同时利用其内置BM25算法实现关键词检索。两种检索结果在ES内部通过rescore阶段进行融合排序，配合三层权限过滤器（userId/orgTag/public），最终返回高质量、可溯源的知识片段。',
    func=[
        '向量KNN检索：基于dense_vector字段存储2048维文本向量，利用ES KNN插件进行近似最近邻搜索',
        '全文BM25检索：利用ES内置BM25算法进行关键词匹配，支持中文分词后的多关键词组合查询',
        '混合检索与rescore重排序：KNN召回宽窗口（recallK=topK×30） + BM25关键词过滤 + rescore阶段加权融合（KNN权重0.2，BM25权重1.0）',
        '三层权限过滤：bool should查询同时匹配userId==requesterId / orgTag∈userOrgTags / isPublic==true',
        '元数据存储：每条文档包含title、content、vector、fileMd5、pageNumber、userId、orgTag、isPublic等字段',
        '批量索引（Bulk API）：文档入库时使用Bulk API批量写入ES索引，大幅提升索引吞吐量',
    ],
    perf='• 知识库索引名：knowledge_base\n• 向量维度：2048（dense_vector类型，使用cosine相似度）\n• KNN召回窗口：recallK = topK × 30（例如topK=5时召回150条候选）\n• 搜索响应时间：<100ms（ES内部检索+rescore）\n• 索引容量：支持百万级文档分块\n• 中文分词：IK Analysis插件（smart模式）\n• ES JVM堆：建议-Xms1g -Xmx1g',
    input=[
        {'name': 'SearchRequest (查询)', 'type': 'ES Query DSL JSON', 'source': 'M1(HybridSearchService)', 'desc': '包含queryVector(float[2048])、keywords(List<String>)、userId、orgTags、topK'},
        {'name': 'BulkIndexRequest (索引)', 'type': 'ES Bulk API JSON', 'source': 'M1(VectorizationService)', 'desc': '包含多条EsDocument(向量+文本+元数据)的批量索引请求'},
    ],
    output=[
        {'name': 'SearchResult列表 (查询结果)', 'type': 'List<SearchResult>', 'source': 'M1(HybridSearchService)', 'desc': '包含chunkId、text、fileName、score、pageNumber、fileMd5的排序列表'},
        {'name': 'BulkResponse (索引确认)', 'type': 'JSON', 'source': 'M1(VectorizationService)', 'desc': '每条文档的索引成功/失败状态'},
    ],
    algo='【ES索引Mapping结构】\n'
         '{\n'
         '  "mappings": {\n'
         '    "properties": {\n'
         '      "chunk_id": {"type": "keyword"},\n'
         '      "title": {"type": "text", "analyzer": "ik_smart"},\n'
         '      "content": {"type": "text", "analyzer": "ik_smart"},\n'
         '      "vector": {"type": "dense_vector", "dims": 2048, "index": true, "similarity": "cosine"},\n'
         '      "file_md5": {"type": "keyword"},\n'
         '      "page_number": {"type": "integer"},\n'
         '      "user_id": {"type": "keyword"},\n'
         '      "org_tag": {"type": "keyword"},\n'
         '      "is_public": {"type": "boolean"}\n'
         '    }\n'
         '  }\n'
         '}\n\n'
         '【混合检索算法（HybridSearchService.searchWithPermission）】\n'
         'Step 1: KNN向量召回\n'
         '  - 参数：field=vector, query_vector=float[2048], k=recallK(topK×30), num_candidates=recallK×2\n'
         '  - ES内部使用HNSW图算法进行近似最近邻搜索\n'
         'Step 2: BM25关键词过滤\n'
         '  - 将查询文本分词（IK smart），构造bool must查询\n'
         '  - 对KNN召回的候选集进行关键词匹配过滤\n'
         'Step 3: Rescore重排序（第二阶段）\n'
         '  - window_size=recallK\n'
         '  - score = KNN_weight × normalize(knn_score) + BM25_weight × normalize(bm25_score)\n'
         '  - 默认权重：KNN=0.2, BM25=1.0（偏向关键词精确匹配）\n'
         'Step 4: 权限过滤器\n'
         '  - bool should: [{term: {user_id: requesterId}}, {term: {org_tag: {value: userOrgTag}}}, {term: {is_public: true}}]\n'
         '  - minimum_should_match=1（至少满足一条）\n'
         'Step 5: 排序取topK，返回结果',
    flow='混合检索作为问答流程的子步骤被M1的HybridSearchService调用，其输入来自M3(查询向量化)，输出返回M1供LLM生成回答。完整流程参见M1模块流程图。',
    iface=[
        {'name': 'HybridSearchService.searchWithPermission()', 'type': 'Java Service', 'method': 'ES Java Client API', 'desc': '构造并执行混合检索DSL，解析ES响应为SearchResult列表'},
        {'name': 'ElasticsearchService.bulkIndex()', 'type': 'Java Service', 'method': 'ES Bulk API', 'desc': '批量索引文档向量，支持部分失败模式'},
        {'name': 'ES REST API (内部)', 'type': 'HTTPS', 'method': 'POST /knowledge_base/_search', 'desc': 'Elasticsearch原生搜索端点（供kibana/调试使用）'},
    ],
    storage='• ES索引文件存储在ES数据目录（docker volume: es_data）\n• 索引mapping和settings持久化在ES集群状态中\n• 单条文档存储：文本内容+向量(2048×4字节=8KB)+元数据，约10KB/chunk',
    constraints=[
        '• 必须运行ES 8.10及以上版本（dense_vector KNN索引需要8.x）',
        '• IK中文分词插件必须提前安装',
        '• 向量维度2048在索引创建时固定，修改需重建索引',
        '• ES为独立JVM进程，需单独分配堆内存（建议1GB）',
        '• 安全性：ES需启用HTTPS和Basic认证（xpack安全功能）',
        '• 索引操作非ACID事务，需在应用层处理一致性（先写MySQL再写ES）',
    ],
    test=[
        {'item': 'KNN向量检索精度', 'method': '人工标注测试集', 'input': '10个查询问题+人工标注的相关文档', 'expected': 'Top-5命中率>80%'},
        {'item': '混合检索vs纯KNN对比', 'method': 'A/B对比测试', 'input': '相同查询在不同模式下的结果', 'expected': '混合检索的NDCG@5优于纯KNN'},
        {'item': '权限过滤器验证', 'method': '不同角色用户搜索同关键词', 'input': 'ADMIN用户 vs 普通用户搜索同一关键词', 'expected': '普通用户看不到非公开文档'},
        {'item': '批量索引性能', 'method': '1000条文档批量写入', 'input': '1000条EsDocument', 'expected': 'Bulk API完成时间<5秒，零错误'},
    ],
)

# --- M5: MySQL ---
reg(7, 'M5', 'MySQL关系数据库',
    desc='M5模块是系统的主数据存储，基于MySQL 8.0，通过JPA/Hibernate进行ORM映射。MySQL存储所有结构化业务数据：用户信息、文件元数据、文档分块（仅文本，不含向量）、对话记录、组织标签、充值记录等。值得注意的是，企业RAG采用"MySQL存文本分块 + ES存向量"的两级存储架构——MySQL中的document_vectors表只存储文本内容，向量数据专属ES存储。',
    func=[
        '用户数据管理：user表存储用户身份、密码哈希、角色、Token余额',
        '文件元数据管理：file_upload表存储上传文件的MD5、名称、大小、向量化状态',
        '文档分块存储：document_vectors表存储文本分块（chunkId、textContent、pageNumber、anchorText），不包含向量',
        '对话记录存储：conversation和chat_message表记录所有问答历史和引用映射',
        '组织标签管理：organization_tag表存储多租户组织树形结构',
        '分片上传记录：chunk_info表记录每个上传分片的元数据（合并后清理）',
        '充值记录：recharge_order表存储用户充值和Token余额变动',
    ],
    perf='• 数据库引擎：InnoDB（支持事务、行锁、外键）\n• 字符集：utf8mb4（支持emoji和生僻字）\n• 连接池：HikariCP（默认20连接，最大50）\n• DDL策略：JPA update模式（开发）/ validate模式（生产）\n• 慢查询阈值：200ms\n• 索引策略：file_md5主键索引、user_id B+树索引、created_at时间索引',
    input=[
        {'name': 'JPA Entity操作', 'type': 'Java对象', 'source': 'M1各Service/Repository', 'desc': 'Spring Data JPA的save/findById/delete等标准CRUD操作'},
    ],
    output=[
        {'name': 'JPA Entity查询结果', 'type': 'Java对象', 'source': 'M1各Service', 'desc': 'findBy...查询、分页查询、聚合统计结果'},
    ],
    algo='【核心表结构设计】\n'
         '1. file_upload表（文件元数据）：\n'
         '   - 主键：file_md5 (VARCHAR(32))\n'
         '   - 关键字段：file_name, file_size, user_id, org_tag, is_public, vectorization_status\n'
         '   - 状态机：UPLOADING → PROCESSING → COMPLETED / FAILED\n\n'
         '2. document_vectors表（文本分块）：\n'
         '   - 主键：id (BIGINT AUTO_INCREMENT)\n'
         '   - 外键：file_md5 REFERENCES file_upload\n'
         '   - 内容字段：chunk_id, text_content, page_number, anchor_text\n'
         '   - 注意：此表<b>不存储向量</b>，向量仅存储在ES的knowledge_base索引中\n\n'
         '3. chat_message表（对话消息）：\n'
         '   - 关联字段：conversation_id REFERENCES conversation\n'
         '   - 内容字段：role(user/assistant/tool), content(LONGTEXT), reference_mappings(JSON)\n'
         '   - reference_mappings示例：{"1": {"fileMd5":"a1b2...", "chunkId":42, "pageNumber":15, "fileName":"差旅制度.pdf"}}',
    flow='MySQL作为持久化存储，几乎每个业务流程都涉及对其的读写。具体读写时机参见M1和M7模块的数据流分析。',
    iface=[
        {'name': 'JpaRepository<User, Long>', 'type': 'Spring Data', 'method': 'ORM自动生成', 'desc': '用户表CRUD + findByUsername自定义查询'},
        {'name': 'JpaRepository<FileUpload, String>', 'type': 'Spring Data', 'method': 'ORM自动生成', 'desc': '文件表CRUD + findByUserId分页查询'},
        {'name': 'JpaRepository<DocumentVector, Long>', 'type': 'Spring Data', 'method': 'ORM自动生成', 'desc': '分块表批量INSERT + findByFileMd5查询'},
        {'name': 'JpaRepository<Conversation, String>', 'type': 'Spring Data', 'method': 'ORM自动生成', 'desc': '对话表CRUD + findByUserId查询'},
    ],
    storage='• MySQL数据文件存储在docker volume: mysql_data\n• 总数据量预估：学生项目千级文档×百级分块=十万级记录，约数百MB\n• JPA自动生成DDL，生产环境建议使用Flyway进行数据库版本管理',
    constraints=[
        '• MySQL为独立进程，需单独管理和备份',
        '• LONGTEXT字段（chat_message.content）单条最大4GB，实际受JVM内存限制',
        '• JSON类型字段（reference_mappings）不支持MySQL 5.7及以下版本',
        '• InnoDB行锁在大批量INSERT时可能产生锁竞争，建议批量操作使用@Transactional',
        '• HikariCP连接池耗尽时请求会阻塞等待，需合理设置连接超时',
    ],
    test=[
        {'item': 'JPA实体映射验证', 'method': 'Spring Boot Test', 'input': '启动时JPA自动建表', 'expected': '所有表结构生成正确，外键约束生效'},
        {'item': '批量INSERT性能', 'method': '1000条DocumentVector写入', 'input': '1000条分块记录', 'expected': 'saveAll完成时间<2秒'},
        {'item': 'JSON字段读写', 'method': '对话记录读写', 'input': '含reference_mappings的消息', 'expected': 'JSON序列化/反序列化正确无损'},
    ],
)

print("Module data definitions part 1 complete. Continuing with M6-M11 + frontend modules...")

# --- M6: Redis ---
reg(8, 'M6', 'Redis缓存服务',
    desc='M6模块是系统的高速缓存层，基于Redis 7.0.11。Redis承担三项核心职责：(1)用户会话和JWT Token缓存，加速认证校验；(2)文件上传进度追踪（Bitmap记录分片完成状态）；(3)API速率限制（滑动窗口计数器）。Redis的大量使用使企业RAG在高并发场景下保持低延迟响应。',
    func=[
        'JWT Token缓存：以tokenId为key存储userId/role/orgTags，加速每次API请求的认证校验',
        'WebSocket会话注册：以userId为key记录活跃WebSocket连接，支持断线检测和重连',
        '文件上传进度追踪：使用Bitmap（SETBIT/GETBIT）记录大文件每个分片的完成状态',
        'API速率限制：基于滑动窗口算法（ZSET），按userId+API类型统计窗口内请求次数',
        '对话缓存：缓存最近20条对话消息，减少MySQL查询',
        '反馈临时存储：用户的赞/踩反馈暂时存储在Redis Hash中，异步批量写入MySQL',
    ],
    perf='• 内存数据库，单次操作延迟<1ms\n• 最大内存：建议maxmemory 512MB\n• 淘汰策略：volatile-lru（仅淘汰设置了过期时间的key）\n• 数据持久化：AOF（appendonly yes, everysec策略）\n• Key命名规范：{业务域}:{userId}:{资源}:{标识符}',
    input=[
        {'name': 'Redis命令', 'type': 'Jedis API调用', 'source': 'M1各Service', 'desc': 'GET/SET/DEL/SETBIT/INCR/ZADD等标准Redis命令'},
    ],
    output=[
        {'name': 'Redis响应', 'type': 'Jedis返回值', 'source': 'M1各Service', 'desc': 'String/Hash/Boolean值，用于判断缓存命中、速率是否超限等'},
    ],
    algo='【速率限制算法（滑动窗口）】\n'
         'Key: rate:{api}:{userId}:{window_start_timestamp}\n'
         '1. 收到请求 → ZREMRANGEBYSCORE key 0 (now - window_size)  # 清理过期记录\n'
         '2. ZCARD key → 获取当前窗口请求数\n'
         '3. 若count >= limit → 拒绝（429 Too Many Requests）\n'
         '4. 若count < limit → ZADD key now_score unique_id → 放行\n'
         '5. EXPIRE key window_size  # 窗口过期自动清理\n\n'
         '【上传进度Bitmap】\n'
         'Key: upload:{userId}:{fileMd5}\n'
         'SETBIT key chunkIndex 1  # 每上传一个分片设置对应位\n'
         'BITCOUNT key → 已上传分片数\n'
         'totalChunks == BITCOUNT → 所有分片已上传完毕，可以触发合并',
    flow='Redis作为缓存层，在认证、上传、限流、对话等多个流程中被M1调用，属于辅助增强角色，不参与核心业务数据流。',
    iface=[
        {'name': 'RedisTemplate / StringRedisTemplate', 'type': 'Spring Data Redis', 'method': 'Jedis连接池', 'desc': 'M1中所有Redis操作的统一入口'},
    ],
    storage='• Redis数据存储在内存中，AOF文件持久化到docker volume: redis_data\n• Key设计规范确保不同业务域的数据逻辑隔离\n• 所有Key均设置合理的过期时间（Token 7天、上传进度 24小时、限流窗口 1分钟）',
    constraints=[
        '• Redis为独立进程，不可用时：认证降级为直接查MySQL、上传进度不可追踪、限流降级为无限制',
        '• 内存容量限制，超maxmemory时按volatile-lru策略淘汰，可能丢失缓存',
        '• AOF持久化非实时（everysec），极端情况下丢失最后1秒数据',
        '• 单机Redis无高可用，生产环境建议Redis Sentinel或Cluster',
    ],
    test=[
        {'item': '缓存读写', 'method': '单元测试', 'input': 'SET key value → GET key', 'expected': 'GET返回value，TTL正确'},
        {'item': '速率限制滑动窗口', 'method': '并发测试', 'input': '1分钟内发送15次请求（限制10次/分钟）', 'expected': '前10次通过，后5次被拒绝'},
        {'item': '上传进度Bitmap', 'method': '模拟分片上传', 'input': '100个分片逐步SETBIT', 'expected': 'BITCOUNT逐步增加，全部完成后为100'},
    ],
)

# --- M7: Kafka ---
reg(9, 'M7', 'Kafka消息队列',
    desc='M7模块是系统的异步消息中间件，基于Apache Kafka 3.2.1。Kafka在企业RAG中承担唯一但关键的职责：解耦文件上传与文档处理。用户上传文件合并完成后，M1通过Kafka事务性发送处理任务消息，立即返回"上传成功"给用户；Kafka Consumer在后台异步消费消息，完成文档解析→切分→向量化→ES索引的完整流水线。这种设计使用户无需等待20-100秒的重计算过程。',
    func=[
        '异步任务分发：将文档处理任务（FileProcessingTask）从上传线程转移到Consumer线程',
        '生产者事务保证：使用KafkaTemplate.executeInTransaction()确保消息不丢失',
        '消费者重试机制：处理失败自动重试3次（间隔3秒），超过重试次数写入DLT死信队列',
        '死信队列（DLT）：保存最终处理失败的任务，保留原始topic/partition/offset信息，支持人工介入',
        '任务类型路由：通过taskType字段区分UPLOAD_PROCESS（新文档）和REINDEX（重索引）',
    ],
    perf='• Topic: file-processing-topic1（单分区，保证顺序消费）\n• 消息格式：JSON序列化的FileProcessingTask\n• 消息大小：约200-500字节/条\n• 消费速度：取决于后端处理速度（20-100秒/条）\n• 消费者并发：单线程顺序消费（保证同一文件的分块处理顺序）',
    input=[
        {'name': 'FileProcessingTask消息', 'type': 'JSON (Kafka消息)', 'source': 'M1(UploadService → KafkaTemplate)', 'desc': '包含fileMd5/filePath/fileName/userId/orgTag/isPublic/taskType'},
    ],
    output=[
        {'name': 'Consumer处理方法触发', 'type': 'Java方法调用', 'source': 'M7 Consumer → M1(FileProcessingConsumer)', 'desc': '触发processTask()，执行完整的异步文档处理流水线'},
    ],
    algo='【消息结构-FileProcessingTask JSON Schema】\n'
         '{\n'
         '  "fileMd5": "a1b2c3d4...",\n'
         '  "filePath": "merged/a1b2c3d4...",\n'
         '  "fileName": "企业制度2025.pdf",\n'
         '  "userId": "user_001",\n'
         '  "orgTag": "default",\n'
         '  "isPublic": false,\n'
         '  "taskType": "UPLOAD_PROCESS | REINDEX",\n'
         '  "requesterId": "user_001"\n'
         '}\n\n'
         '【生产者-消费者流程】\n'
         'Producer端（UploadController.mergeFile()）：\n'
         '  1. MinIO composeObject合并分片成功\n'
         '  2. MySQL INSERT file_upload记录\n'
         '  3. kafkaTemplate.executeInTransaction(template -> template.send(topic, task))\n'
         '  4. 返回用户 {"status":"success","message":"文件已提交处理"}\n\n'
         'Consumer端（FileProcessingConsumer.processTask()）：\n'
         '  1. 从M8下载合并后的文件\n'
         '  2. 调用M9(LiteParse)/M10(Tika)解析文档\n'
         '  3. 调用M11(HanLP)分词 → 滑动窗口切分\n'
         '  4. 分块写入M5(MySQL) document_vectors表\n'
         '  5. 分块通过M3(Embedding)向量化\n'
         '  6. 向量写入M4(ES) knowledge_base索引\n'
         '  7. 更新M5(MySQL) file_upload状态为COMPLETED\n\n'
         '【失败处理】\n'
         '  - SeekToCurrentErrorHandler: 重试3次（间隔3秒）\n'
         '  - DeadLetterPublishingRecoverer: 3次后写入DLT死信topic\n'
         '  - DLT消息保留原始topic/partition/offset头信息',
    flow='管理员上传文档的完整异步处理流程。PhaseA（上传/合并/分发）同步完成并立即响应用户；PhaseB（解析/切分/向量化/索引）在Kafka Consumer中异步执行：',
    flow_img='03_upload_flow.png',
    iface=[
        {'name': 'KafkaTemplate.send(topic, task)', 'type': 'Spring Kafka', 'method': '事务性发送', 'desc': '生产者发送FileProcessingTask到指定topic'},
        {'name': '@KafkaListener(topics="file-processing-topic1")', 'type': 'Spring Kafka', 'method': '消费者注解', 'desc': '自动拉取消息并调用FileProcessingConsumer.processTask()'},
    ],
    storage='• Kafka消息存储在docker volume: kafka_data\n• 消息保留策略：7天或1GB（先到达为准）\n• DLT死信消息永久保留，直到人工处理',
    constraints=[
        '• Kafka为独立进程，需单独部署和维护（含Zookeeper/KRaft）',
        '• 单分区Topic保证顺序消费，但限制了并行处理能力（同一时间只能处理一个文件）',
        '• 消息无事务回滚：若Consumer处理到一半崩溃，MySQL分块可能已写入但ES索引未完成（应用层最终一致性）',
        '• 学生项目文件处理量小（<100文件/天），Kafka的分布式能力实际上是过度设计',
        '• Kafka不可用时：文档上传的异步处理完全不工作，文件上传成功但永远不会被向量化',
    ],
    test=[
        {'item': '消息发送与消费', 'method': '集成测试', 'input': '发送一个FileProcessingTask', 'expected': 'Consumer在10秒内接收并开始处理'},
        {'item': '事务性发送验证', 'method': '模拟Kafka故障', 'input': '发送过程中断开Kafka连接', 'expected': '事务回滚，M1向用户返回上传失败'},
        {'item': '消费重试与DLT', 'method': '模拟处理失败', 'input': 'Consumer处理时抛出异常', 'expected': '重试3次后消息进入DLT，原始topic/offset信息保留'},
        {'item': '消息幂等性', 'method': '发送重复消息', 'input': '同一fileMd5发送两次', 'expected': 'Consumer检测到重复，跳过或覆盖处理'},
    ],
)

# --- M8: MinIO ---
reg(10, 'M8', 'MinIO对象存储',
    desc='M8模块是系统的文件存储层，基于MinIO 8.5.12（AWS S3兼容对象存储）。MinIO负责存储所有用户上传的原始文件，支持大文件的分片暂存和服务端合并。M8通过S3兼容SDK与M1通信，提供预签名URL供Consumer下载文件，支持公开URL转换供前端预览。',
    func=[
        '分片存储：将大文件的每个5MB分片以独立对象存储在chunks/{fileMd5}/{chunkIndex}路径',
        '服务端合并：使用composeObject API在MinIO服务端将多个分片合并为完整文件，无需下载再上传',
        '文件下载：Consumer通过预签名URL（1小时有效期）安全下载合并后的文件',
        '文件删除：管理员删除文档时，同时删除merged/{fileMd5}的原始文件',
        '公开URL转换：将内部S3 URL转换为可公开访问的HTTP URL，供前端文件预览使用',
    ],
    perf='• Bucket: uploads\n• 单个分片大小：5MB（与前端约定一致）\n• composeObject合并速度：约100MB/s（MinIO服务端操作，不经过网络）\n• 预签名URL有效期：1小时\n• 存储容量：受限于MinIO所在磁盘（docker volume: minio_data）',
    input=[
        {'name': '文件分片数据', 'type': '二进制流 (multipart)', 'source': 'M1(UploadController)', 'desc': '每片5MB的文件分片二进制数据'},
        {'name': '合并指令', 'type': 'composeObject API调用', 'source': 'M1(UploadService)', 'desc': '源对象列表 chunks/{md5}/0,1,2... → 目标对象 merged/{md5}'},
    ],
    output=[
        {'name': '对象存储URL', 'type': 'S3对象路径', 'source': 'M1 → M7(Consumer)', 'desc': 's3://uploads/merged/{fileMd5}，Consumer通过预签名URL下载'},
        {'name': '公开访问URL', 'type': 'HTTP URL', 'source': 'M1 → 前端', 'desc': 'http://minio:9000/uploads/merged/{fileMd5}（公开可读）'},
    ],
    algo='【存储路径设计】\n'
         'uploads/\n'
         '├── chunks/{fileMd5}/{chunkIndex}  ← 上传中分片（合并后清理）\n'
         '└── merged/{fileMd5}               ← 合并完成的完整文件（长期存储）\n\n'
         '【composeObject合并过程】\n'
         '1. 前端调用POST /api/v1/upload/merge\n'
         '2. M1构造composeObject请求：\n'
         '   sourceObjects: [chunks/{md5}/0, chunks/{md5}/1, ..., chunks/{md5}/N-1]\n'
         '   destinationObject: merged/{md5}\n'
         '3. MinIO服务端直接在存储层合并（不经过M1应用内存）\n'
         '4. 合并成功后：\n'
         '   - 删除所有chunks/{md5}/分片（MinIO removeObjects）\n'
         '   - 清理M6(Redis)上传进度Bitmap\n'
         '   - 清理M5(MySQL) chunk_info分片记录\n'
         '5. 构造FileProcessingTask发送到M7(Kafka)',
    flow='M8在文档上传流程中承担分片暂存和文件合并的职责。详见M7模块的文档上传流程图。',
    flow_img='03_upload_flow.png',
    iface=[
        {'name': 'MinioClient.putObject()', 'type': 'MinIO Java SDK', 'method': 'S3兼容API', 'desc': '向指定bucket和路径写入对象（分片）'},
        {'name': 'MinioClient.composeObject()', 'type': 'MinIO Java SDK', 'method': 'S3兼容API', 'desc': '服务端合并多个对象为一个（核心上传完成操作）'},
        {'name': 'MinioClient.getPresignedObjectUrl()', 'type': 'MinIO Java SDK', 'method': 'S3兼容API', 'desc': '生成有时效的预签名下载URL'},
        {'name': 'MinioClient.removeObject()', 'type': 'MinIO Java SDK', 'method': 'S3兼容API', 'desc': '删除指定对象'},
    ],
    storage='• MinIO数据存储在docker volume: minio_data\n• 数据格式：原始文件二进制（PDF/DOCX/TXT等，无特殊编码）\n• 存储空间预估：学生项目几百个PDF，每个平均5MB，总计约数GB',
    constraints=[
        '• MinIO为独立进程，不可用时无法上传新文件',
        '• composeObject限制：单次最多合并1000个源对象（即最多1000个分片，单文件最大5GB）',
        '• 预签名URL有效期内可被任何人访问（Bearer URL机制），需注意链接泄露风险',
        '• 单机MinIO无高可用，生产环境建议MinIO集群或云S3服务',
        '• MinIO默认Access Key/Secret Key认证，生产环境需定期轮换',
    ],
    test=[
        {'item': '分片上传与合并', 'method': '集成测试', 'input': '上传10个5MB分片后触发合并', 'expected': '合并后merged/{md5}文件大小为50MB，MD5校验一致'},
        {'item': 'composeObject原子性', 'method': '合并过程中断测试', 'input': '合并到一半kill MinIO进程', 'expected': '合并失败，chunks保留，merged未生成'},
        {'item': '预签名URL下载', 'method': 'HTTP GET', 'input': '使用生成的预签名URL', 'expected': '1小时内可直接下载，超时后403 Forbidden'},
    ],
)

# --- M9: LiteParse ---
reg(11, 'M9', 'LiteParse PDF解析引擎',
    desc='M9模块专门负责PDF文件的解析，是一个本地命令行工具（CLI，命令名为lit）。它能够提取PDF中每页的文本内容，保留页码信息，并支持OCR识别扫描件中的文字。M9输出结构化的JSON格式（按页组织），M1的ParseService解析此JSON后交给M11进行分词切块。M9与M10(Tika)分工明确：M9处理PDF，M10处理其余所有格式。',
    func=[
        'PDF文本提取：解析PDF内部文本流，输出每页的纯文本内容',
        '页码保留：输出JSON中按页组织，保留page_number字段，用于引用溯源',
        'OCR文字识别：对扫描版PDF（图片型）启用OCR引擎，支持中英文混合识别',
        '段落级文本块：输出blocks数组，每个block包含文本内容和坐标位置',
    ],
    perf='• 解析速度：普通PDF约2-10秒/100页，OCR模式约30秒/100页\n• 最大页数：--max-pages参数限制（默认1000页）\n• OCR语言：chi_sim+eng（中文简体+英文）\n• OCR DPI：150（平衡速度与精度）\n• CLI启动开销：约500ms（Java进程fork）',
    input=[
        {'name': 'PDF文件路径', 'type': '本地文件系统路径', 'source': 'M1(ParseService, 从M8下载后的临时文件)', 'desc': '已下载到本地的PDF文件绝对路径'},
    ],
    output=[
        {'name': '解析结果JSON', 'type': 'JSON文件（写入/tmp/parsed/）', 'source': 'M1(ParseService)', 'desc': '{"pages":[{"page_number":1,"text":"...","blocks":[...]},...]}'},
    ],
    algo='【CLI命令格式】\n'
         'lit parse --input "merged/a1b2c3.pdf" \\\n'
         '           --output-dir "/tmp/parsed/" \\\n'
         '           --ocr --lang chi_sim+eng \\\n'
         '           --dpi 150 --max-pages 1000\n\n'
         '【M1调用流程（ParseService.parsePdfWithLiteParse）】\n'
         '1. 构建CLI命令参数\n'
         '2. ProcessBuilder.start() 启动子进程\n'
         '3. 等待进程结束（process.waitFor(120, TimeUnit.SECONDS)）\n'
         '4. 读取输出JSON文件\n'
         '5. 使用Jackson反序列化为LiteParseResult Java对象\n'
         '6. 遍历pages，提取每页text和page_number\n'
         '7. 将文本传递给M11(HanLP)进行分词和切块',
    flow='M9在Kafka Consumer的异步文档处理流程中被调用（仅针对PDF文件）。详细流程参见M7模块流程图。',
    iface=[
        {'name': 'lit parse (CLI命令)', 'type': '命令行接口', 'method': 'ProcessBuilder fork子进程', 'desc': 'PDF解析命令，输入文件路径，输出结构化的JSON解析结果'},
    ],
    storage='• LiteParse本身不存储数据，解析结果JSON写入/tmp/parsed/临时目录\n• 解析完成后M1读取JSON并删除临时文件\n• 文本分块持久化到M5(MySQL) document_vectors表',
    constraints=[
        '• LiteParse必须预先安装为系统CLI工具（lit命令在PATH中）',
        '• OCR功能依赖语言包（chi_sim+eng），需提前下载',
        '• 子进程fork开销：每个PDF启动一个lit进程，大文件需较长超时时间（120秒）',
        '• CLI调用非线程安全：并行处理多个PDF时需确保每个任务使用不同的临时输出目录',
        '• LiteParse为Python工具，需Python环境和依赖库',
    ],
    test=[
        {'item': '普通PDF解析', 'method': 'CLI调测', 'input': '一个10页的中文PDF', 'expected': '输出JSON包含10个pages，每页text字段非空，page_number正确'},
        {'item': 'OCR扫描件解析', 'method': 'CLI调测(带--ocr)', 'input': '一个图片型PDF', 'expected': '提取到文字内容（准确率>85%）'},
        {'item': '大文件解析超时', 'method': '边界测试', 'input': '一个500页的PDF', 'expected': '120秒内返回结果或超时抛异常'},
    ],
)

# --- M10: Apache Tika ---
reg(12, 'M10', 'Apache Tika通用文档解析引擎',
    desc='M10模块是系统的通用文档解析引擎，基于Apache Tika的AutoDetectParser。与M9(LiteParse)专注PDF不同，M10负责解析Word(.doc/.docx)、TXT、Markdown、HTML等非PDF格式文档。Tika自动检测文件MIME类型并选择合适的解析器，输出统一格式的纯文本字符串，供后续分词和切块使用。Tika以纯Java库形式集成在M1中，无需独立进程。',
    func=[
        '文件类型自动检测：通过文件头魔数（Magic Bytes）自动识别文档格式',
        'Word文档解析：支持.doc（OLE格式）和.docx（OOXML格式）的文本提取',
        '纯文本文档读取：TXT/Markdown/HTML等格式的文本提取和HTML标签清理',
        '元数据提取：提取文档的作者、标题、创建时间等元信息（当前企业RAG未充分利用）',
    ],
    perf='• 非IO密集型操作：解析速度取决于文件大小，通常<1秒/页\n• 内存占用：Tika解析器为流式处理，不将整个文件加载到内存\n• 支持的格式：30+种文档格式（通过AutoDetect自动分发）\n• 纯Java实现：无外部进程开销，比M9(LiteParse CLI)快5-10倍',
    input=[
        {'name': '文件输入流', 'type': 'java.io.InputStream', 'source': 'M1(ParseService, 从M8下载的文件流)', 'desc': '任意格式文档（Word/TXT/MD/HTML等）的字节流'},
    ],
    output=[
        {'name': '纯文本内容', 'type': 'String', 'source': 'M1(ParseService) → M11(HanLP) → M5(MySQL)', 'desc': '已剥离格式的纯文本字符串'},
    ],
    algo='【调用方式（ParseService.parseFile）】\n'
         '1. 通过MIME类型（或文件扩展名）判断文件格式\n'
         '2. 若为PDF → 委托给M9(LiteParse CLI)处理\n'
         '3. 若为其他格式 → 使用Tika AutoDetectParser：\n'
         '   - Metadata metadata = new Metadata()\n'
         '   - metadata.set(Metadata.CONTENT_TYPE, mimeType)\n'
         '   - BodyContentHandler handler = new BodyContentHandler(-1)  // 无长度限制\n'
         '   - parser.parse(inputStream, handler, metadata, new ParseContext())\n'
         '   - String text = handler.toString()\n'
         '4. 返回纯文本，进入M11(HanLP)分词切块阶段',
    flow='M10在Kafka Consumer的异步文档处理流程中被调用（非PDF格式文件）。处理入口为M1的ParseService.parseFile()方法。详细流程参见M7模块流程图。',
    iface=[
        {'name': 'AutoDetectParser.parse()', 'type': 'Java API (Tika库)', 'method': '同步方法调用', 'desc': '自动检测文件类型并提取文本内容到ContentHandler'},
    ],
    storage='• Tika本身不存储数据，纯文本输出传递给M11和M5\n• Tika以Maven依赖形式嵌入M1（pom.xml: org.apache.tika:tika-core + tika-parsers-standard-package）',
    constraints=[
        '• Tika解析.docx依赖POI库，某些复杂格式（如嵌入OLE对象）可能解析不完整',
        '• 对于扫描型PDF（图片PDF），Tika默认只提取嵌入文本，不会OCR（需委托给M9）',
        '• BodyContentHandler(-1)无输出大小限制，理论上可能OOM（恶意超大文件需在Controller层限制）',
        '• Tika的AutoDetectParser不是线程安全的，每次解析需创建新实例',
    ],
    test=[
        {'item': 'Word(.docx)解析', 'method': '单元测试', 'input': '一个包含文字、表格、图片的docx文件', 'expected': '提取到完整的文字内容，格式标记被移除'},
        {'item': 'Markdown解析', 'method': '单元测试', 'input': '一个含代码块的Markdown文件', 'expected': '提取到纯文本，Markdown语法标记被移除'},
        {'item': '未知格式处理', 'method': '边界测试', 'input': '一个二进制非文档文件', 'expected': 'Tika抛出TikaException，M1捕获后标记解析失败'},
    ],
)

# --- M11: HanLP ---
reg(13, 'M11', 'HanLP中文分词引擎',
    desc='M11模块是系统的中文自然语言处理基础组件，基于HanLP的StandardTokenizer。M11在文档分块阶段被调用，对M9/M10解析出的纯文本进行中文分词，确保滑动窗口切分的边界尽可能落在完整的词语边界上（而非生硬地按字符数截断），从而提高检索召回率。M11以纯Java库形式集成，无外部进程。',
    func=[
        '中文分词：使用StandardTokenizer对中文文本进行词语切分（基于CRF模型）',
        '词性标注：标注每个词语的词性（名词/动词/形容词等，当前企业RAG仅用分词结果）',
        '分词引导切块：基于分词边界优化滑动窗口切分，避免在词语中间截断',
    ],
    perf='• 分词速度：约100万字符/秒（单线程）\n• 内存占用：约200MB（首次加载HanLP模型）\n• 分词精度：在通用中文文本上准确率>95%\n• 纯Java实现：无外部进程开销',
    input=[
        {'name': '纯文本字符串', 'type': 'String', 'source': 'M1(ParseService, 来自M9或M10的解析输出)', 'desc': '待分词的原始文本'},
    ],
    output=[
        {'name': '分词序列', 'type': 'List<Term>', 'source': 'M1(ParseService)', 'desc': '每个词及其词性（如: [企业/n, 知识/n, 库/n, RAG/nx, 问答/vn, 系统/n]）'},
    ],
    algo='【分词+切块算法】\n'
         '1. 接收解析后的纯文本 + 参数：chunk_size=512 tokens, overlap=100 tokens, min_chunk_size=100\n'
         '2. 调用HanLP StandardTokenizer.segment(text) → List<Term>\n'
         '3. 按token边界滑动窗口切分：\n'
         '   - 指针start=0\n'
         '   - 循环：\n'
         '     a. 从start开始累积tokens至总数达到chunk_size\n'
         '     b. 寻找最近的句号/换行符作为切分点（优先保证语义完整）\n'
         '     c. 若无天然断点，则在当前完整token后切分\n'
         '     d. 生成一个DocumentVector（chunkId, textContent, pageNumber, anchorText）\n'
         '     e. start = 当前切分点 - overlap（滑动窗口重叠）\n'
         '   - 当start >= totalTokens时结束\n'
         '4. 返回List<DocumentVector>（每条含文本内容+元数据，不含向量）',
    flow='M11在Kafka Consumer的异步处理流程中的分词切块步骤被调用。输入来自M9/M10的输出，输出传递给M5(MySQL存储)和M3(向量化)。详细流程参见M7模块流程图。',
    iface=[
        {'name': 'HanLP.newSegment() / StandardTokenizer', 'type': 'Java API', 'method': '同步方法调用', 'desc': '创建分词器实例，调用segment()进行分词'},
    ],
    storage='• HanLP本身不存储数据，以Maven依赖形式嵌入M1（pom.xml: com.hankcs:hanlp）\n• 模型文件在首次使用时自动下载到本地缓存目录（~/.hanlp/）\n• 分词结果不单独存储，仅用于指导切块（切块结果存入M5 document_vectors表）',
    constraints=[
        '• HanLP模型首次加载需下载约200MB的模型文件（需要网络或预置模型包）',
        '• StandardTokenizer基于CRF序列标注，对于高度专业化的企业术语（如"差旅标准"）可能切分不准确',
        '• 分词速度虽快，但在处理超长文档（>10万字符）时会产生明显的停顿',
        '• HanLP以LGPL协议开源，商业使用需注意许可合规',
    ],
    test=[
        {'item': '基本分词准确率', 'method': '单元测试+人工评估', 'input': '"企业知识库RAG问答系统是一个基于检索增强生成技术的系统"', 'expected': '"企业/知识库/RAG/问答/系统"被正确切分为独立词条'},
        {'item': '切块边界优化验证', 'method': '检查切块边界', 'input': '一段含句号的自然段落', 'expected': '切块边界优先落在句号/段落结束处，而非在词语中间截断'},
        {'item': 'overlap正确性', 'method': '检查相邻块内容', 'input': '切块参数chunk_size=512, overlap=100', 'expected': '相邻两个chunk的最后100个token内容重叠'},
    ],
)

print("Backend modules M1-M11 defined. Continuing with M-F01 to M-F11 frontend modules...")

# ═══════════════════ FRONTEND MODULES ═══════════════════

# --- M-F05: Auth Module ---
reg(14, 'M-F05', '认证模块（前端）',
    desc='M-F05是前端的用户认证模块，负责登录、注册、Token管理和用户信息缓存。该模块基于Vue 3 Composition API和Pinia状态管理实现，与后端/auth/*接口对接。用户登录成功后获得JWT Token，存储在localStorage中，每次API请求由M-F04自动注入Authorization头。',
    func=[
        '密码登录：用户名+密码表单提交，获取JWT Token和用户信息',
        '用户注册：含邀请码验证的注册表单，支持表单校验和API错误提示',
        'Token管理：Token本地持久化（localStorage），支持自动刷新和过期处理',
        '用户信息缓存：登录成功后缓存用户信息到Pinia authStore，供全局使用',
        '登出：清除Token和用户信息，跳转到登录页',
        '角色判断：提供isAdmin/isUser等getter，支持前端权限控制',
    ],
    perf='• 登录响应时间：<200ms（同步HTTP请求）\n• Token本地存储：localStorage（持久化，关闭浏览器后保留）\n• 认证状态响应：Pinia响应式，状态变更立即反映到UI',
    input=[
        {'name': '登录表单数据', 'type': '{username, password}', 'source': '登录页面表单', 'desc': '用户输入的登录凭证'},
        {'name': '注册表单数据', 'type': '{username, password, inviteCode, ...}', 'source': '注册页面表单', 'desc': '新用户注册信息'},
    ],
    output=[
        {'name': 'JWT Token', 'type': 'String (Bearer Token)', 'source': 'localStorage + Pinia authStore', 'desc': '用于API请求认证的令牌'},
        {'name': '用户信息', 'type': 'Api.UserInfo', 'source': 'Pinia authStore', 'desc': '{userId, username, role, orgTags, tokenBalance}'},
    ],
    algo='【登录流程】\n'
         '1. 用户填写表单 → 点击登录按钮\n'
         '2. 前端调用POST /auth/login {username, password}\n'
         '3. 后端返回 {code:0, data: {token, userInfo}}\n'
         '4. authStore.setToken(token) → 写入localStorage\n'
         '5. authStore.setUserInfo(userInfo) → 更新Pinia状态\n'
         '6. router.push("/chat") → 跳转到问答页\n\n'
         '【Token自动刷新流程】\n'
         '1. M-F04请求拦截器检测到401响应\n'
         '2. 调用POST /auth/refresh携带旧Token\n'
         '3. 若刷新成功 → 更新Token → 重试原请求\n'
         '4. 若刷新失败 → 清除Token → 弹出登录框',
    flow='登录流程为标准的表单提交→验证→存储→跳转流程，无复杂分支逻辑。',
    iface=[
        {'name': 'POST /auth/login', 'type': 'REST API', 'method': 'HTTP POST', 'desc': '用户登录，返回JWT Token和用户信息'},
        {'name': 'POST /auth/register', 'type': 'REST API', 'method': 'HTTP POST', 'desc': '用户注册，需携带邀请码'},
        {'name': 'GET /auth/userinfo', 'type': 'REST API', 'method': 'HTTP GET', 'desc': '获取当前用户信息'},
        {'name': 'POST /auth/refresh', 'type': 'REST API', 'method': 'HTTP POST', 'desc': '刷新即将过期的Token'},
    ],
    storage='• JWT Token存储：localStorage（key: "token"），持久化保存\n• 用户信息：Pinia authStore内存状态 + localStorage备份\n• 登出时：清除localStorage + 重置Pinia store',
    constraints=[
        '• JWT Token有过期时间（默认7天），过期后需重新登录',
        '• localStorage中的Token可被XSS攻击窃取，需配合CSP策略使用',
        '• 注册功能依赖邀请码模块（若后端简化可改为开放注册）',
    ],
    test=[
        {'item': '正常登录', 'method': 'E2E测试', 'input': '正确的用户名和密码', 'expected': '跳转到/chat页面，authStore.userInfo非空'},
        {'item': '错误密码', 'method': 'E2E测试', 'input': '错误的密码', 'expected': '显示"用户名或密码错误"，停留在登录页'},
        {'item': 'Token过期自动处理', 'method': '模拟过期Token', 'input': '发起API请求时Token已过期', 'expected': '弹出重新登录对话框'},
    ],
)

# --- M-F06: Knowledge Base ---
reg(15, 'M-F06', '知识库管理模块（前端）',
    desc='M-F06是前端的知识库管理模块，提供文档上传、文件列表展示、删除和向量化状态查询功能。该模块支持大文件分片上传（每片5MB + MD5校验 + 断点续传），上传进度实时展示，以及文件预览（PDF/图片/文本）。知识库管理面向所有登录用户，但部分操作（如查看所有人文件）仅限管理员。',
    func=[
        '文件列表展示：分页表格显示已上传文档（文件名、大小、上传时间、向量化状态、嵌入Token消耗）',
        '大文件分片上传：前端计算MD5、切割为5MB分片、逐片上传、实时进度展示',
        '断点续传：上传中断后可恢复（通过MD5检测已上传分片）',
        '文件删除：确认对话框后删除文档（软删除，MySQL标记）',
        '文件预览：通过API获取文件内容并在模态框中预览（支持PDF/图片/文本）',
        '向量化状态：实时展示文档的向量化处理状态（UPLOADING→PROCESSING→COMPLETED/FAILED）',
        '权限设置：上传时可选择公开/私有、关联组织标签',
    ],
    perf='• 分片上传并发：单文件分片串行上传（保证顺序），多文件可并行\n• 分片大小：5MB（与后端约定一致）\n• MD5计算：SparkMD5分块计算，不阻塞UI\n• 文件列表加载：分页查询，每页20条',
    input=[
        {'name': '用户选择的文件', 'type': 'File对象 (HTML5 File API)', 'source': '浏览器文件选择器', 'desc': '支持.pdf/.doc/.docx/.txt格式，单文件最大100MB'},
        {'name': '上传配置参数', 'type': '{orgTag, isPublic}', 'source': '上传对话框表单', 'desc': '组织标签和公开/私有设置'},
    ],
    output=[
        {'name': '上传进度', 'type': '百分比 + 分片状态', 'source': '上传对话框进度条', 'desc': '实时显示已上传分片数/总分片数、上传速度'},
        {'name': '文件列表', 'type': 'Api.FileUpload[]', 'source': '知识库页面表格', 'desc': '分页文档列表，含向量化状态'},
    ],
    algo='【分片上传算法】\n'
         '1. 用户选择文件 → 前端SparkMD5计算完整文件MD5（分块增量计算，不阻塞UI）\n'
         '2. 查询后端是否已存在同一MD5的上传任务（断点续传检测）\n'
         '3. 文件分为ceil(fileSize/5MB)个分片\n'
         '4. 每个分片：POST /upload/chunk {file, fileMd5, chunkIndex, totalChunks}\n'
         '5. 进度计算：(已上传分片数/总分片数) × 100%\n'
         '6. 所有分片上传完成 → POST /upload/merge {fileMd5, fileName, totalChunks}\n'
         '7. 合并成功 → 刷新文件列表 → 显示"文件已提交处理"',
    flow='上传流程为前端驱动的分阶段操作：MD5计算→分片上传→进度追踪→合并触发→列表刷新。详见M7模块的文档上传流程图。',
    flow_img='03_upload_flow.png',
    iface=[
        {'name': 'POST /api/v1/upload/chunk', 'type': 'REST API', 'method': 'HTTP POST (multipart)', 'desc': '上传单个5MB分片'},
        {'name': 'POST /api/v1/upload/merge', 'type': 'REST API', 'method': 'HTTP POST (JSON)', 'desc': '触发分片合并和Kafka任务分发'},
        {'name': 'GET /api/v1/documents', 'type': 'REST API', 'method': 'HTTP GET', 'desc': '分页获取文档列表'},
        {'name': 'DELETE /api/v1/documents/{fileMd5}', 'type': 'REST API', 'method': 'HTTP DELETE', 'desc': '删除文档及关联数据'},
    ],
    storage='• 上传分片临时存储：前端File对象分片（Blob），不持久化\n• MD5计算结果暂存：组件内存状态\n• 文件列表数据：Pinia kbStore + Naive UI DataTable',
    constraints=[
        '• 浏览器File API限制：单文件最大受浏览器内存限制（Chrome约2GB）',
        '• MD5计算耗时：Large文件（100MB）约需3-5秒（SparkMD5增量计算）',
        '• 分片上传为串行模式（保证服务端顺序），多文件并行需多个上传实例',
    ],
    test=[
        {'item': '小文件上传', 'method': 'E2E测试', 'input': '一个1MB的PDF文件', 'expected': '上传成功，列表中出现该文件，状态为PROCESSING→COMPLETED'},
        {'item': '大文件分片上传', 'method': 'E2E测试', 'input': '一个50MB的PDF文件', 'expected': '分片进度条正确展示10个分片完成情况，合并后列表显示'},
        {'item': '断点续传', 'method': 'E2E测试', 'input': '上传到50%时断网→恢复网络→重新上传同一文件', 'expected': '检测到已有分片，从第6个分片继续上传'},
    ],
)

# --- M-F07: Chat (Q&A) ---
reg(16, 'M-F07', '智能问答模块（前端）',
    desc='M-F07是前端最核心、最复杂的业务模块，实现基于WebSocket的实时流式问答交互。该模块负责对话管理、消息渲染（Markdown+代码高亮）、引用溯源展示、用户反馈收集以及WebSocket连接生命周期管理。模块采用Vue 3 Composition API + Pinia状态管理 + @vueuse/core useWebSocket技术栈，是RAG管道在前端的最终呈现。',
    func=[
        '对话管理：新建对话、切换历史对话、对话归档/取消归档',
        'WebSocket流式问答：建立WebSocket连接，发送用户问题，实时接收AI的逐token回答',
        'Markdown渲染：将AI生成的Markdown文本渲染为富文本（表格、代码块高亮、列表等）',
        '引用溯源展示：在AI回答中标注来源编号（来源#1: filename），点击可查看原文片段',
        '工具调用可视化：展示ReAct Agent的工具调用过程（"正在检索知识库..." → "检索完成"）',
        '用户反馈：回答后的赞/踩按钮，调用submit_feedback工具',
        '停止生成：用户可随时停止正在生成的回答',
        '速率限制提示：在输入框旁显示剩余请求次数和重置时间',
    ],
    perf='• WebSocket连接建立：<500ms（含JWT认证）\n• 首token展示延迟：<200ms\n• 流式渲染性能：使用requestAnimationFrame批量更新DOM，避免频繁渲染\n• 消息列表虚拟滚动：支持1000+条消息不卡顿',
    input=[
        {'name': '用户问题文本', 'type': 'String', 'source': '聊天输入框', 'desc': '用户输入的自然语言问题'},
        {'name': '停止生成信号', 'type': 'JSON帧 {type:"stop"}', 'source': '用户点击停止按钮', 'desc': '中断当前正在生成的LLM回答'},
        {'name': '反馈数据', 'type': '{rating: "up"/"down", messageId}', 'source': '用户点击赞/踩按钮', 'desc': '对AI回答质量的评价'},
    ],
    output=[
        {'name': '流式AI回答', 'type': 'WebSocket JSON帧序列', 'source': '聊天消息区域', 'desc': '打字机效果展示AI回答，中间穿插工具调用状态提示'},
        {'name': '引用来源链接', 'type': '可点击的引用标记', 'source': 'AI回答中嵌入', 'desc': '"来源#1: 差旅制度2025.pdf"，点击跳转到原文预览'},
    ],
    algo='【WebSocket通信协议】\n'
         '前端发送（JSON帧）：\n'
         '  {type: "ask", conversationId: "uuid", question: "什么是RAG？"}\n'
         '  {type: "stop", generationId: "gen_xxx"}\n'
         '  {type: "ping"}  // 心跳保活\n\n'
         '后端推送（JSON帧序列）：\n'
         '  {type: "start", generationId: "gen_xxx"}\n'
         '  {type: "chunk", chunk: "RAG"}  // 逐token文本\n'
         '  {type: "tool_call", tool: "search_knowledge", status: "running"}\n'
         '  {type: "tool_call", tool: "search_knowledge", status: "completed"}\n'
         '  {type: "completion", status: "finished", generationId: "gen_xxx",\n'
         '   referenceMappings: {"1": {"fileMd5":"...","chunkId":42,"fileName":"差旅制度2025.pdf"}}}\n\n'
         '【消息渲染流程】\n'
         '1. 收到chunk帧 → 追加到当前AI消息的content字符串末尾\n'
         '2. 使用markdown-it将Markdown转为HTML（含highlight.js代码高亮）\n'
         '3. DOMPurify过滤XSS风险 → 设置v-html渲染\n'
         '4. 收到completion帧 → 标记消息完成，显示引用链接和反馈按钮\n'
         '5. 调用POST /conversations/{id}/messages保存完整对话',
    flow='用户问答交互流程：输入问题 → WebSocket send → 逐token接收 → 实时渲染 → 完成展示引用。完整数据流参见M1模块的ReAct Agent流程图。',
    flow_img='02_qa_flow.png',
    iface=[
        {'name': 'WebSocket ws://host/chat/{token}', 'type': 'WebSocket (浏览器原生)', 'method': 'new WebSocket(url) / @vueuse/core useWebSocket', 'desc': '建立双向长连接，支持文本帧收发'},
        {'name': 'GET /api/v1/conversations', 'type': 'REST API', 'method': 'HTTP GET', 'desc': '获取用户对话列表'},
        {'name': 'POST /api/v1/conversations', 'type': 'REST API', 'method': 'HTTP POST', 'desc': '创建新对话'},
    ],
    storage='• 对话列表缓存：Pinia chatStore + localStorage备份\n• 当前消息列表：组件内存状态，切换对话时从API重新加载\n• WebSocket实例：@vueuse/core useWebSocket管理，组件卸载时自动关闭',
    constraints=[
        '• WebSocket连接受浏览器同源策略和网络代理限制',
        '• 用户刷新页面 → WebSocket断开 → 需重新建立连接和加载对话历史',
        '• 切换对话时旧WebSocket连接中的生成任务需先停止',
        '• Markdown渲染使用v-html，必须配合DOMPurify防止XSS攻击',
        '• 移动端浏览器可能限制后台WebSocket连接（微信小程序不支持原生WebSocket）',
    ],
    test=[
        {'item': '基本问答流程', 'method': 'E2E测试', 'input': '"企业2025年差旅标准是什么？"', 'expected': 'WebSocket连接成功，逐token显示回答，最后显示引用来源链接'},
        {'item': '停止生成功能', 'method': 'E2E测试', 'input': '生成过程中点击停止按钮', 'expected': '回答停止，显示"已停止生成"，对话记录保留已生成部分'},
        {'item': '对话切换', 'method': 'E2E测试', 'input': '在对话A和对话B之间切换', 'expected': '消息列表正确切换，各自保留独立的消息记录'},
        {'item': 'Markdown渲染', 'method': '组件测试', 'input': 'AI返回含表格、代码块、列表的回答', 'expected': '表格正确渲染、代码语法高亮、列表格式正确'},
    ],
)

# --- M-F01 to M-F04: Infrastructure modules (brief) ---
for seq, mid, name, desc_data in [
    (17, 'M-F01', '路由管理模块（前端）', {
        'desc': 'M-F01基于Vue Router 4和@elegant-router/vue插件实现。@elegant-router/vue通过扫描views/目录的文件结构自动生成路由定义，支持懒加载、路由守卫（认证检查、角色校验、页面标题更新）和三种history模式。',
        'func': ['自动路由生成：扫描views/目录生成16个路由的定义和懒加载映射',
                 '路由守卫链：认证守卫（未登录→登录页）、角色守卫（非管理员→403）、标题守卫（更新document.title）',
                 '多模式支持：hash模式（开发）/ history模式（生产需服务端配置）/ memory模式'],
        'perf': '• 路由懒加载：每个页面独立chunk，首屏仅加载当前页面代码\n• 路由切换动画：CSS transition 300ms平滑过渡',
        'iface': [{'name': 'router/index.ts', 'type': 'Vue Router实例', 'method': 'createRouter()', 'desc': '创建路由实例，注册守卫和插件'}],
        'constraints': ['• hash模式下URL中带有#，SEO不友好（管理后台无需SEO）', '• history模式需要服务端配置fallback到index.html'],
    }),
    (18, 'M-F02', '布局框架模块（前端）', {
        'desc': 'M-F02基于Soybean Admin的AdminLayout组件，提供管理后台的完整页面骨架。包括：顶部导航栏（Logo+菜单切换+用户头像）、可折叠侧边栏、多标签页栏、面包屑导航、内容区域（RouterView+keep-alive缓存）和页脚。',
        'func': ['5种布局模式：垂直菜单、水平菜单、垂直混合、水平混合、反向水平混合',
                 '多标签页管理：打开/关闭/切换页签、右键菜单（关闭当前/关闭其他/关闭左侧/关闭所有）',
                 '侧边栏：一级/二级菜单展开折叠、固定/浮动模式切换',
                 '内容区缓存：keep-alive基于路由名称缓存页面状态'],
        'perf': '• 布局切换：响应式断点（xs/sm/md/lg/xl），移动端自动折叠侧边栏\n• 标签页缓存：最多缓存10个页面实例',
        'constraints': ['• 布局组件为Soybean Admin框架提供，修改需谨慎避免破坏响应式行为', '• 标签页缓存占用内存，页面数超过10个时自动移除最早未访问的缓存'],
    }),
    (19, 'M-F03', '主题配置模块（前端）', {
        'desc': 'M-F03管理前端的视觉主题配置，基于Naive UI的NConfigProvider实现。支持亮/暗/自动三种模式、自定义主题色、布局模式切换和页面功能开关。配置通过Pinia themeStore持久化到localStorage。',
        'func': ['亮/暗/自动模式：跟随系统偏好自动切换（prefers-color-scheme媒体查询）',
                 '主题色自定义：预设多套配色方案，实时预览切换效果',
                 '页面功能开关：侧边栏折叠、页眉显示、标签页显示、页脚显示等可分别控制'],
        'perf': '• 主题切换：CSS变量注入模式，切换<100ms无需重新渲染',
        'constraints': ['• 主题色需与Naive UI组件兼容，自定义色值需通过色彩算法生成调色板', '• 暗模式下部分自定义组件（如PDF查看器）需额外适配'],
    }),
    (20, 'M-F04', '请求基础设施模块（前端）', {
        'desc': 'M-F04是前端与后端通信的基础设施，基于@sa/axios封装。提供统一的请求/响应拦截器：请求拦截器自动注入JWT Token到Authorization头；响应拦截器统一处理业务错误码、401未认证（弹出登录框）和网络异常。',
        'func': ['JWT自动注入：每个请求自动附加Authorization: Bearer {token}',
                 '401拦截处理：检测到401响应 → 尝试刷新Token → 失败则弹出登录模态框',
                 '统一错误处理：网络异常、超时、服务端错误的通用提示',
                 '请求重试：GET请求网络异常时自动重试1次'],
        'perf': '• 请求超时：默认30秒（上传接口为0不限时）\n• 拦截器性能：同步操作<1ms，不影响请求延迟',
        'iface': [{'name': 'createFlatRequest()', 'type': '@sa/axios封装', 'method': '函数调用', 'desc': '创建配置好的axios实例，导出为全局request对象'}],
        'constraints': ['• 依赖后端返回统一的{code, message, data}响应格式', '• Token刷新接口需幂等，避免并发刷新导致多次请求'],
    }),
]:
    reg(seq, mid, name, **desc_data)

# --- M-F08 to M-F11: Additional business modules ---
reg(21, 'M-F08', '用户管理模块（前端）',
    desc='M-F08是管理员专用的用户管理页面，提供用户列表查询、角色分配、组织标签设置和Token配额管理功能。普通用户无法访问此页面（路由守卫角色检查）。',
    func=['用户列表：分页表格显示所有用户（用户名、角色、组织标签、Token余额、注册时间）',
          '用户搜索：按用户名或组织标签过滤',
          '角色分配：切换用户角色（USER/ADMIN）',
          '组织标签分配：为用户关联/解绑组织标签（级联选择器）',
          'Token配额追加：管理员为用户添加Token额度（弹出对话框输入数量）'],
    perf='• 用户列表分页：每页20条，支持排序和筛选\n• 搜索防抖：输入后300ms触发请求',
    iface=[{'name': 'GET/PUT /api/v1/users', 'type': 'REST API', 'method': 'HTTP GET/PUT', 'desc': '用户列表查询、用户信息修改'}],
)

reg(22, 'M-F09', '混合搜索模块（前端）',
    desc='M-F09是知识库中的混合搜索功能，以对话框形式提供。用户输入关键词后，前端调用POST /search/hybrid接口，后端执行KNN向量检索+BM25关键词检索的混合搜索，返回排序后的结果列表。搜索结果的每条记录包含文件名、匹配片段和相似度评分。',
    func=['混合搜索对话框：关键词输入 → API调用 → 结果列表展示',
          '搜索结果展示：显示匹配片段（含高亮关键词）、文件名、评分',
          '文件预览跳转：点击搜索结果可跳转到对应文档的预览页面'],
    iface=[{'name': 'POST /api/v1/search/hybrid', 'type': 'REST API', 'method': 'HTTP POST', 'desc': '执行KNN+BM25混合搜索，返回排序结果'}],
)

reg(23, 'M-F10', '模型配置模块（前端）',
    desc='M-F10是管理员专用的LLM/Embedding Provider配置页面。管理员可以管理多个模型提供商（如DeepSeek、DashScope）的API Key和端点URL，设置当前活跃的Provider，以及测试API连接是否正常。',
    func=['Provider列表管理：CRUD操作管理模型服务商',
          'API Key配置：安全存储和显示API Key（脱敏显示）',
          '活跃Provider切换：选择当前使用的LLM和Embedding Provider',
          '连接测试：发送测试请求验证API Key和端点的有效性'],
    iface=[{'name': 'GET/POST/PUT/DELETE /api/v1/model-providers', 'type': 'REST API', 'method': 'HTTP CRUD', 'desc': '模型Provider的完整CRUD操作'}],
)

reg(24, 'M-F11', '使用监控模块（前端）',
    desc='M-F11是管理员专用的系统使用监控仪表盘，基于ECharts实现数据可视化。展示Token消耗趋势图、API调用频率、用户用量排行榜、速率限制配置表单和异常告警列表。',
    func=['Token消耗图表：ECharts折线图展示每日/每小时Token消耗趋势',
          '调用排行榜：按用户展示API调用次数和Token消耗排名',
          '速率限制配置：管理员动态调整各API的速率限制参数',
          '异常告警列表：展示近期的速率限制触发记录和异常调用'],
    iface=[{'name': 'GET /api/v1/usage/*', 'type': 'REST API', 'method': 'HTTP GET', 'desc': '获取各类使用统计和监控数据'}],
)

print(f"Total modules defined: {len(MODULES)} (11 backend + 13 frontend/infrastructure)")

# ═══════════════════ DOCUMENT GENERATION ═══════════════════

# ── Cover Page ──
DOC.add_paragraph()
DOC.add_paragraph()
DOC.add_paragraph()
p = DOC.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('企业知识库RAG问答系统')
run.bold = True; run.font.size = Pt(26); run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = DOC.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('详细设计说明书')
run.bold = True; run.font.size = Pt(22); run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

DOC.add_paragraph(); DOC.add_paragraph()
P('基于 企业RAG 完整原型架构（含全部外部服务）', size=12)
P('文档版本：V1.0', size=12)
P('生成日期：2026-06-24', size=12)
DOC.add_page_break()

# ── 1. 引言 ──
DOC.add_heading('1. 引言', level=1)

DOC.add_heading('1.1 编写目的', level=2)
P('本文档的编制目的是对企业知识库RAG问答系统进行详细的模块级设计，明确各模块的功能、接口、处理流程、数据结构和测试方案，为后续的编码实现、单元测试和系统集成提供技术依据。本文档的预期读者包括：')
P('• 软件开发人员：依据本文档进行模块编码和单元测试')
P('• 系统设计人员：验证设计方案的一致性和完整性')
P('• 项目管理人员：跟踪设计进度和模块完成情况')
P('• 测试人员：依据接口定义编写集成测试用例')
P('• 指导教师/评审专家：评估系统设计的合理性和完备性')

DOC.add_heading('1.2 背景', level=2)
T(['项目属性', '说明'], [
    ['系统名称', '企业知识库RAG问答系统（企业RAG / 派聪明）'],
    ['任务提出者', '软件工程课程设计/毕业设计选题'],
    ['开发者', '项目团队（3-5人）'],
    ['目标用户', '企业员工、客服人员、技术支持人员、管理人员'],
    ['后端技术栈', 'Spring Boot 3.4.2 + Java 17 + Elasticsearch 8.10 + Kafka 3.2 + MySQL 8.0 + Redis 7.0 + MinIO 8.5'],
    ['前端技术栈', 'Vue 3.5 + TypeScript 5.8 + Vite 6.3 + Naive UI 2.41 + Pinia 3.0'],
    ['AI服务', 'DeepSeek API (LLM) + DashScope text-embedding-v4 (Embedding)'],
    ['参考原型', '企业RAG v1.3.13 (https://github.com/AsRe6666/企业RAG)'],
    ['开发周期', '3周（需求分析1周 + 详细设计+编码1.5周 + 测试0.5周）'],
], [3, 12])

DOC.add_heading('1.3 定义', level=2)
T(['术语/缩写', '定义'], [
    ['RAG', 'Retrieval-Augmented Generation，检索增强生成——在LLM生成回答前，先从知识库中检索相关文档片段作为上下文，提高回答的准确性和可溯源性'],
    ['ReAct', 'Reasoning + Acting，一种Agent模式，LLM交替进行推理（Reasoning）和工具调用（Acting），自主决定何时检索、何时生成'],
    ['Embedding', '文本向量化——将文本转换为固定维度的浮点数向量，使语义相近的文本在向量空间中距离相近'],
    ['KNN', 'K-Nearest Neighbors，K近邻搜索——在向量空间中查找与查询向量最相似的K个向量'],
    ['BM25', 'Best Matching 25，一种基于词频和逆文档频率的关键词检索评分算法，Elasticsearch默认使用'],
    ['RRF', 'Reciprocal Rank Fusion，倒数排名融合——将多个排序列表合并为一个统一排序的算法'],
    ['SSE', 'Server-Sent Events，服务器推送事件——HTTP协议下的单向流式数据传输'],
    ['JWT', 'JSON Web Token，一种用于身份认证的加密令牌，包含用户身份和权限信息'],
    ['DLT', 'Dead Letter Topic，死信队列——Kafka中存储最终处理失败消息的专用topic，供人工介入排查'],
    ['OCR', 'Optical Character Recognition，光学字符识别——从扫描件/图片型PDF中提取文字的技术'],
    ['WebSocket', '双向全双工通信协议，用于实时流式问答，比HTTP轮询更高效'],
    ['MoE', 'Mixture of Experts，混合专家模型架构——DeepSeek使用的大模型架构'],
], [4, 11])

DOC.add_heading('1.4 参考资料', level=2)
P('• 企业RAG 项目源码：https://github.com/AsRe6666/企业RAG')
P('• Soybean Admin 前端框架文档：https://docs.soybeanjs.com')
P('• Spring Boot 3.4 官方文档：https://docs.spring.io/spring-boot/')
P('• Elasticsearch 8.10 文档：https://www.elastic.co/guide/en/elasticsearch/reference/8.10/')
P('• Apache Kafka 3.2 文档：https://kafka.apache.org/documentation/')
P('• DeepSeek API 文档：https://platform.deepseek.com/api-docs')
P('• DashScope Embedding 文档：https://help.aliyun.com/document_detail/dashscope.html')
P('• HanLP 文档：https://github.com/hankcs/HanLP')
P('• Naive UI 组件文档：https://www.naiveui.com/')
P('• 选题分析报告：report.md')
P('• 前端架构分析文档：frontend.md')

DOC.add_page_break()

# ── 2. 系统的结构 ──
DOC.add_heading('2. 系统的结构', level=1)
P('本系统由11个后端软件模块和11个前端模块组成。后端模块按照"软件/服务"进行划分，每个模块对应一个独立的软件组件或外部服务。模块间通过HTTP REST API、JDBC、TCP Socket、CLI调用、Java API等方式进行通信。系统架构总览如下图所示：')
IMG('01_system_architecture.png', 5.5)

DOC.add_heading('2.1 模块标识符与层次关系', level=2)
P('系统模块按层次分为五层：')
T(['层次', '编号', '标识符', '模块名称', '软件/服务', '类型'],
  [
    ['用户层', '—', 'Browser', 'Web浏览器', 'Vue 3 + Naive UI 前端SPA', '客户端'],
    ['应用编排层', '1', 'M1', 'Spring Boot应用服务框架', 'Spring Boot 3.4.2', '核心框架'],
    ['AI服务层', '2', 'M2', 'DeepSeek API大语言模型', 'DeepSeek API (云端)', '外部API'],
    ['AI服务层', '3', 'M3', 'DashScope Embedding向量化', 'DashScope text-embedding-v4 (云端)', '外部API'],
    ['中间件层', '4', 'M4', 'Elasticsearch检索引擎', 'Elasticsearch 8.10.0', '中间件服务'],
    ['中间件层', '5', 'M5', 'MySQL关系数据库', 'MySQL 8.0 (InnoDB)', '中间件服务'],
    ['中间件层', '6', 'M6', 'Redis缓存服务', 'Redis 7.0.11', '中间件服务'],
    ['中间件层', '7', 'M7', 'Kafka消息队列', 'Apache Kafka 3.2.1', '中间件服务'],
    ['中间件层', '8', 'M8', 'MinIO对象存储', 'MinIO 8.5.12 (S3兼容)', '中间件服务'],
    ['解析层', '9', 'M9', 'LiteParse PDF解析引擎', 'LiteParse CLI (Python)', '本地工具'],
    ['解析层', '10', 'M10', 'Apache Tika通用文档解析', 'Apache Tika 2.x', 'Java库'],
    ['解析层', '11', 'M11', 'HanLP中文分词引擎', 'HanLP Portable', 'Java库'],
  ], [2.5, 1.5, 1.5, 4.5, 4, 1.5])

DOC.add_heading('2.2 模块依赖关系矩阵', level=2)
P('下表展示各后端模块之间的依赖关系（✅ = 行模块直接调用/读写列模块）：')
T(['行\\列', 'M1', 'M2', 'M3', 'M4', 'M5', 'M6', 'M7', 'M8', 'M9', 'M10', 'M11'],
  [
    ['M1', '—', '✅', '✅', '✅', '✅', '✅', '✅', '✅', '✅', '✅', '✅'],
    ['M2', '', '—', '', '', '', '', '', '', '', '', ''],
    ['M3', '', '', '—', '', '', '', '', '', '', '', ''],
    ['M4', '', '', '', '—', '', '', '', '', '', '', ''],
    ['M5', '', '', '', '', '—', '', '', '', '', '', ''],
    ['M6', '', '', '', '', '—', '', '', '', '', '', ''],
    ['M7(Consumer)', '回调M1', '', '回调M3', '回调M4', '回调M5', '', '—', '回调M8', '回调M9', '回调M10', '回调M11'],
  ], [3, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2])
P('注：M2-M6是纯数据/计算服务，不主动调用其他模块。M7(Kafka)的Consumer运行在M1进程内，通过M1的Service层回调所有其他模块。')

DOC.add_heading('2.3 系统部署架构', level=2)
IMG('06_deployment.png', 5.0)

DOC.add_heading('2.4 数据库ER图', level=2)
P('本系统使用MySQL作为主数据库，核心表结构及关系如下图所示：')
IMG('05_er_diagram.png', 5.0)

DOC.add_heading('2.5 前端模块体系', level=2)
P('前端基于Soybean Admin脚手架，分为基础设施模块（框架提供，仅需配置）和业务模块（自主开发）。共11个模块：')
IMG('04_frontend_architecture.png', 5.0)
T(['编号', '模块名称', '类型', '说明'],
  [
    ['M-F01', '路由管理模块', '基础设施', '页面路由定义、懒加载、三种路由守卫（认证/角色/标题）'],
    ['M-F02', '布局框架模块', '基础设施', '管理后台骨架：5种布局模式、多标签页、面包屑'],
    ['M-F03', '主题配置模块', '基础设施', '亮/暗/自动模式、主题色、CSS变量注入'],
    ['M-F04', '请求基础设施模块', '基础设施', 'axios封装、JWT自动注入、401拦截、错误统一处理'],
    ['M-F05', '认证模块', '业务模块', '登录/注册/Token管理/用户信息缓存/角色判断'],
    ['M-F06', '知识库管理模块', '业务模块', '文档分片上传/MD5校验/断点续传/列表/预览/删除'],
    ['M-F07', '智能问答模块', '业务模块', 'WebSocket流式问答/Markdown渲染/引用溯源/反馈/停止生成'],
    ['M-F08', '用户管理模块', '业务模块', '用户列表/角色分配/组织标签/Token配额管理'],
    ['M-F09', '混合搜索模块', '业务模块', '知识库KNN+BM25混合检索/结果高亮/文件预览跳转'],
    ['M-F10', '模型配置模块', '业务模块', 'LLM/Embedding Provider CRUD/API Key管理/连接测试'],
    ['M-F11', '使用监控模块', '业务模块', 'ECharts折线图/Token消耗趋势/排行榜/速率限制配置'],
  ], [2, 4, 2.5, 6.5])

DOC.add_page_break()

# ── Generate All Module Sections ──
print("Generating all module sections...")
for seq, mid, name, data in MODULES:
    print(f"  Writing module {mid}: {name}")
    module(seq, mid, name, data)

# ── Save ──
output_path = '/home/susan/MySpace/codeSpace/shixun/详细设计说明书_企业RAG完整版.docx'
DOC.save(output_path)
print(f"\nDocument saved to: {output_path}")
print("Done!")
